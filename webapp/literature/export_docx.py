from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from . import LiteratureRecord
from .normalize import build_citation, normalize_text, source_display_label


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 10.5,
    center: bool = False,
) -> None:
    """写入单元格，支持 \\n 换行。"""
    cell.text = ""
    lines = str(text or "").split("\n")
    if not lines:
        return
    p0 = cell.paragraphs[0]
    if center:
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p0.add_run(lines[0])
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    for line in lines[1:]:
        p = cell.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"


def export_records_to_docx(records: list[LiteratureRecord]) -> tuple[bytes, str]:
    """按 Clinical Literature Search Result.docx 结构导出：Database / Item / Literature，并保留 Link。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # 对齐参考模板列宽（单位：EMU）
    col_widths = (763905, 539750, 2910205, 1409700)
    for i, w in enumerate(col_widths):
        table.columns[i].width = w
    headers = ("Database", "Item", "Literature", "Link")
    for i, h in enumerate(headers):
        _set_cell_text(
            table.rows[0].cells[i],
            h,
            bold=True,
            size=10.5,
            center=i in (0, 1),
        )

    # 按来源分组连续编号（与样例文档一致：各库各自 [1]..[n]）
    counters: dict[str, int] = {}
    for rec in records:
        # database 已由 normalize_record 归一化（如 scholar→Google）；直接用，
        # 避免对已映射值再套 source_display_label 造成 Google→GOOGLE 与列表不一致
        db = normalize_text(rec.get("database")) or source_display_label(rec.get("source"))
        counters[db] = counters.get(db, 0) + 1
        item = f"[{counters[db]}]"
        # 强制用已清洗字段重建引用，再 normalize_text 兜底剥 HTML（导出绝不信任脏 citation）
        citation = normalize_text(build_citation(rec)) or normalize_text(rec.get("citation"))
        link = normalize_text(rec.get("source_url"))
        row = table.add_row().cells
        _set_cell_text(row[0], db, size=10.5, center=True)
        _set_cell_text(row[1], item, size=10.5, center=True)
        _set_cell_text(row[2], citation, size=10.5)
        _set_cell_text(row[3], link, size=10.5)
        if link:
            for run in row[3].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    buf = io.BytesIO()
    doc.save(buf)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return buf.getvalue(), f"Clinical_Literature_Search_Result_{ts}.docx"
