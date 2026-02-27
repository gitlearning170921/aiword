"""
文档模板解析与生成相关的服务函数。
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List
from urllib.request import urlopen
from urllib.error import URLError, HTTPError

from docx import Document


def download_template_from_url(url: str, save_path: str) -> str:
    """
    从 URL 下载模板文件到本地并返回保存路径。
    :param url: 模板文件 URL（需可直接 GET 下载）
    :param save_path: 本地保存路径（.docx）
    :return: save_path
    """
    path = Path(save_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        with urlopen(url, timeout=30) as resp:
            data = resp.read()
        path.write_bytes(data)
        return str(path)
    except (URLError, HTTPError, OSError) as e:
        raise RuntimeError(f"下载模板失败：{e}") from e


def _replace_placeholders_in_paragraph(paragraph, mapping: Dict[str, str]):
    if not mapping:
        return

    original_text = paragraph.text
    if not original_text:
        return

    combined_text = "".join(run.text for run in paragraph.runs if run.text) or original_text

    replaced_text = combined_text
    for key, value in mapping.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in replaced_text:
            replaced_text = replaced_text.replace(placeholder, str(value))

    if replaced_text != combined_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = replaced_text
        else:
            paragraph.add_run(replaced_text)


def _replace_placeholders_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, mapping)


def extract_placeholders(template_path: str) -> List[str]:
    """从模板文件中提取占位符列表。"""
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    doc = Document(str(template_path))
    placeholders: List[str] = []
    seen = set()
    pattern = re.compile(r"\{\{\s*([^{}\n\r]+?)\s*\}\}")

    def collect_from_text(text: str):
        if not text or not text.strip():
            return
        if "\t" in text and any(char.isdigit() for char in text):
            return
        for match in pattern.finditer(text):
            placeholder = match.group(1).strip()
            if placeholder and placeholder not in seen and "{" not in placeholder and "}" not in placeholder:
                seen.add(placeholder)
                placeholders.append(placeholder)

    def collect_from_paragraph(paragraph):
        combined = "".join(run.text for run in paragraph.runs if run.text) or paragraph.text
        collect_from_text(combined)

    def collect_from_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_from_paragraph(paragraph)
                for nested_table in cell.tables:
                    collect_from_table(nested_table)

    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    paragraph_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    body = doc.element.body
    for element in body:
        if isinstance(element, CT_P):
            paragraph = paragraph_map.get(element)
            if paragraph:
                collect_from_paragraph(paragraph)
        elif isinstance(element, CT_Tbl):
            table = table_map.get(element)
            if table:
                collect_from_table(table)

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                collect_from_paragraph(paragraph)
            for table in section.header.tables:
                collect_from_table(table)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                collect_from_paragraph(paragraph)
            for table in section.footer.tables:
                collect_from_table(table)

    return placeholders


def generate_document(template_path: str, output_dir: str, data: Dict[str, str], output_name: str | None = None) -> str:
    """根据模板与数据生成 Word 文档，返回生成文件路径。"""
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)

    if not output_name:
        output_name = f"{template_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"

    output_path = output_dir_path / output_name

    doc = Document(str(template_path))

    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, data)

    for table in doc.tables:
        _replace_placeholders_in_table(table, data)

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, data)
            for table in section.header.tables:
                _replace_placeholders_in_table(table, data)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, data)
            for table in section.footer.tables:
                _replace_placeholders_in_table(table, data)

    doc.save(str(output_path))
    return str(output_path)

