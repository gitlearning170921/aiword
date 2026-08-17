"""从 YY-IW-020 制度原文解析版本任务归档表。

支持：
- 本地 .docx（python-docx 按表抽取）
- 知识库重建文本（训练时写入的 Word 表格 TSV 摘录）

匹配最新版：文件名/文号含 YY-IW-020 或「医疗软件质量合规管理制度」，
优先封面「文件版次」，其次文件名中的 Vx.y，再次入库时间。
"""

from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Any, Optional

CHAPTER_CHANGE = "软件变更管理"
CHAPTER_DEFECT = "缺陷管理"
CHAPTER_RELEASE = "软件生产/发布管理"
CHAPTER_TRACE = "系统追溯"

DEFAULT_HEADING_BITS = {
    CHAPTER_CHANGE: ["X", "Y"],
    CHAPTER_TRACE: ["X", "Y"],
    CHAPTER_DEFECT: ["X", "Y", "Z"],
    CHAPTER_RELEASE: ["X", "Y", "Z", "B"],
}

_FILE_NEEDLES = (
    "iw-020",
    "iw020",
    "医疗软件质量合规管理",
)

_VER_IN_NAME = re.compile(r"[Vv](\d+)(?:\.(\d+))?(?:\.(\d+))?")
_VER_IN_TEXT = re.compile(r"文件版次\s*[:：]?\s*([Vv]?\d+(?:\.\d+){0,3})")
_HEADING_BITS = re.compile(r"版本号\s*([XYZBxyzb/／、和与及\s位]+)")
_DOC_NO = re.compile(r"YY\s*-?\s*IW\s*-?\s*020", re.I)


def normalize_rule_filename(name: str) -> str:
    return re.sub(r"\s+", "", (name or "")).casefold().replace("－", "-")


def is_yy_iw_020_filename(name: str) -> bool:
    n = normalize_rule_filename(name)
    if "iw-020" in n or "iw020" in n:
        return True
    return "医疗软件质量合规管理" in (name or "")


def parse_version_tuple(raw: str) -> tuple[int, int, int]:
    text = (raw or "").strip()
    m = _VER_IN_NAME.search(text)
    if not m:
        return (0, 0, 0)
    return (int(m.group(1) or 0), int(m.group(2) or 0), int(m.group(3) or 0))


def extract_version_label(*texts: str) -> str:
    for text in texts:
        s = (text or "").strip()
        if not s:
            continue
        m = _VER_IN_TEXT.search(s)
        if m:
            lab = m.group(1).strip()
            return lab if lab.upper().startswith("V") else f"V{lab}"
        m = _VER_IN_NAME.search(s)
        if m:
            parts = [m.group(1)]
            if m.group(2) is not None:
                parts.append(m.group(2))
            if m.group(3) is not None:
                parts.append(m.group(3))
            return "V" + ".".join(parts)
    return ""


def parse_heading_bits(text: str) -> list[str]:
    m = _HEADING_BITS.search(text or "")
    if not m:
        return []
    found: list[str] = []
    for ch in m.group(1).upper():
        if ch in "XYZB" and ch not in found:
            found.append(ch)
    return found


def _cell(text: Any) -> str:
    return " ".join(str(text or "").split())


def _uniquify_seq(rows: list[tuple]) -> list[tuple]:
    counts = Counter(str(r[0]) for r in rows)
    used: dict[str, int] = {}
    out: list[tuple] = []
    for row in rows:
        seq = str(row[0])
        if counts[seq] > 1:
            used[seq] = used.get(seq, 0) + 1
            seq = f"{seq}{chr(ord('a') + used[seq] - 1)}"
        out.append((seq,) + tuple(row[1:]))
    return out


def _col_index(header: list[str], *needles: str) -> int:
    for i, h in enumerate(header):
        if any(n in (h or "") for n in needles):
            return i
    return -1


def _rows_from_header_grid(header: list[str], data_rows: list[list[str]]) -> list[tuple]:
    joined = "".join(header)
    if "归档物名称" not in joined or "归档频率" not in joined:
        return []
    i_seq = _col_index(header, "序号")
    i_name = _col_index(header, "归档物名称")
    i_freq = _col_index(header, "归档频率")
    i_author = _col_index(header, "编制")
    i_caution = _col_index(header, "注意事项")
    i_content = _col_index(header, "内容是否需要修改")
    if i_name < 0 or i_freq < 0:
        return []
    out: list[tuple] = []
    for raw in data_rows:
        name = _cell(raw[i_name] if i_name < len(raw) else "")
        if not name or name in {"归档物名称", "/"}:
            continue
        seq = _cell(raw[i_seq] if 0 <= i_seq < len(raw) else "")
        freq = _cell(raw[i_freq] if i_freq < len(raw) else "")
        author = _cell(raw[i_author] if 0 <= i_author < len(raw) else "")
        caution = _cell(raw[i_caution] if 0 <= i_caution < len(raw) else "")
        content = _cell(raw[i_content] if 0 <= i_content < len(raw) else "")
        out.append((seq, name, freq, author, caution, content))
    return _uniquify_seq(out)


def _classify_archive_table(
    rows: list[tuple],
    *,
    chapter: str = "",
    sublabel: str = "",
    change_seen: int = 0,
) -> Optional[str]:
    names = {str(r[1]) for r in rows}
    if not names:
        return None
    if "软件发布验证方案" in names or "发布计划" in names or "产品放行申请单" in names:
        return "release"
    if "缺陷评估报告" in names and "缺陷记录表" in names and "软件开发计划" not in names:
        return "defect"
    if "变更申请单" in names and "软件需求规范" in names:
        if "用户接口需求规范" in names or sublabel == "ce" or "CE获证" in chapter:
            return "changeCe"
        if sublabel == "nmpa" or "国内获证" in chapter:
            return "changeNmpa"
        if chapter == CHAPTER_CHANGE or "软件变更" in chapter:
            return "changeCe" if change_seen == 0 else "changeNmpa"
        return "changeCe" if change_seen == 0 else "changeNmpa"
    return None


def _update_heading_state(text: str, state: dict[str, Any]) -> None:
    t = _cell(text)
    if not t:
        return
    bits = parse_heading_bits(t)
    if "软件变更管理" in t:
        state["chapter"] = CHAPTER_CHANGE
        if bits:
            state["headingBits"][CHAPTER_CHANGE] = bits
    elif "系统追溯" in t:
        state["chapter"] = CHAPTER_TRACE
        if bits:
            state["headingBits"][CHAPTER_TRACE] = bits
    elif "缺陷管理" in t and "软件变更" not in t:
        state["chapter"] = CHAPTER_DEFECT
        if bits:
            state["headingBits"][CHAPTER_DEFECT] = bits
    elif "软件生产" in t or "发布管理" in t:
        state["chapter"] = CHAPTER_RELEASE
        if bits:
            state["headingBits"][CHAPTER_RELEASE] = bits
    if t.startswith("CE获证") or t == "CE获证产品":
        state["sublabel"] = "ce"
    elif t.startswith("国内获证") or t == "国内获证产品":
        state["sublabel"] = "nmpa"


def _empty_row_buckets() -> dict[str, list[tuple]]:
    return {
        "changeCe": [],
        "changeNmpa": [],
        "defect": [],
        "release": [],
    }


def _collect_process_hints(text: str) -> dict[str, Any]:
    t = text or ""
    immediate = [
        "变更申请单",
        "软件需求规范",
        "架构设计规范",
        "系统测试方案",
        "软件发布说明",
        "发布记录",
        "检验记录",
    ]
    return {
        "immediateFiles": immediate,
        "releaseSignoff": ("软件发布说明" in t and "签字" in t),
        "postReleaseVerify": ("发布成功后" in t and "测试验证" in t),
    }


def _default_trace_row() -> list[tuple]:
    return [
        (
            "trace",
            "软件可追溯性分析报告",
            "版本号X/Y位变更时",
            "产品经理",
            "日常软件变更需说明需求、开发、测试对应关系，通过禅道管理或者填写本报告",
            "",
        )
    ]


def _payload(
    *,
    rows: dict[str, list[tuple]],
    heading_bits: dict[str, list[str]],
    source_file: str,
    source_version: str,
    source_date: str = "",
    matched_by: str,
    process_hints: Optional[dict[str, Any]] = None,
    candidates: Optional[list[dict[str, Any]]] = None,
    message: str = "",
) -> dict[str, Any]:
    bits = dict(DEFAULT_HEADING_BITS)
    bits.update({k: v for k, v in (heading_bits or {}).items() if v})
    serial_rows = {
        key: [list(item) for item in (rows.get(key) or [])]
        for key in ("changeCe", "changeNmpa", "defect", "release")
    }
    serial_rows["traceability"] = [list(item) for item in _default_trace_row()]
    ok = any(serial_rows[k] for k in ("changeCe", "changeNmpa", "defect", "release"))
    ver = source_version or "V?"
    src = source_file or "YY-IW-020"
    return {
        "ok": ok,
        "matchedBy": matched_by,
        "sourceFile": source_file,
        "sourceVersion": ver,
        "sourceDate": source_date,
        "ruleSource": f"{'知识库' if matched_by == 'knowledge_base' else '本地制度'}《{src}》（{ver}）",
        "headingBits": bits,
        "rows": serial_rows,
        "processHints": process_hints or _collect_process_hints(""),
        "candidates": candidates or [],
        "message": message or ("已解析归档表" if ok else "未解析到归档表"),
    }


def extract_from_docx(path: str | Path, *, matched_by: str = "local_docx") -> dict[str, Any]:
    path = Path(path)
    if not path.is_file():
        return _payload(
            rows=_empty_row_buckets(),
            heading_bits={},
            source_file=path.name,
            source_version="",
            matched_by=matched_by,
            message="制度文件不存在",
        )
    try:
        from docx import Document
    except Exception as exc:
        return _payload(
            rows=_empty_row_buckets(),
            heading_bits={},
            source_file=path.name,
            source_version="",
            matched_by=matched_by,
            message=f"python-docx 不可用：{exc}",
        )
    try:
        doc = Document(str(path))
    except Exception as exc:
        return _payload(
            rows=_empty_row_buckets(),
            heading_bits={},
            source_file=path.name,
            source_version="",
            matched_by=matched_by,
            message=f"无法打开制度文件：{exc}",
        )

    cover_text = ""
    source_date = ""
    if doc.tables:
        cover = [[_cell(c.text) for c in r.cells] for r in doc.tables[0].rows]
        cover_text = " ".join(x for row in cover for x in row)
        for row in cover:
            if row and "日期" in row[0] and len(row) > 1:
                source_date = row[1]
            if len(row) >= 4 and "日期" in row[2]:
                source_date = source_date or row[3]

    state: dict[str, Any] = {
        "chapter": "",
        "sublabel": "",
        "headingBits": {},
    }
    buckets = _empty_row_buckets()
    change_seen = 0
    para_texts: list[str] = []
    tbl_i = 0
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            texts = [n.text or "" for n in child.iter() if n.text]
            t = _cell("".join(texts))
            if t:
                para_texts.append(t)
                _update_heading_state(t, state)
        elif tag == "tbl":
            if tbl_i >= len(doc.tables):
                continue
            table = doc.tables[tbl_i]
            tbl_i += 1
            if not table.rows:
                continue
            header = [_cell(c.text) for c in table.rows[0].cells]
            data = [[_cell(c.text) for c in r.cells] for r in table.rows[1:]]
            rows = _rows_from_header_grid(header, data)
            if not rows:
                continue
            kind = _classify_archive_table(
                rows,
                chapter=str(state.get("chapter") or ""),
                sublabel=str(state.get("sublabel") or ""),
                change_seen=change_seen,
            )
            if kind in {"changeCe", "changeNmpa"}:
                change_seen += 1
            if kind and not buckets[kind]:
                buckets[kind] = rows

    source_version = extract_version_label(cover_text, path.name)
    hints = _collect_process_hints("\n".join(para_texts))
    return _payload(
        rows=buckets,
        heading_bits=state.get("headingBits") or {},
        source_file=path.name,
        source_version=source_version,
        source_date=source_date,
        matched_by=matched_by,
        process_hints=hints,
        message="已从制度原文解析归档表" if any(buckets.values()) else "制度原文未解析到归档表",
    )


def _parse_tsv_block(block: str) -> list[tuple]:
    lines = [ln.rstrip("\n") for ln in (block or "").splitlines() if ln.strip()]
    header_idx = -1
    header: list[str] = []
    for i, ln in enumerate(lines):
        cols = [c.strip() for c in ln.split("\t")]
        if "归档物名称" in "".join(cols) and "归档频率" in "".join(cols):
            header_idx = i
            header = cols
            break
    if header_idx < 0:
        return []
    data = [[c.strip() for c in ln.split("\t")] for ln in lines[header_idx + 1 :]]
    return _rows_from_header_grid(header, data)


def extract_from_text(
    text: str,
    *,
    source_file: str = "",
    matched_by: str = "knowledge_base",
    source_date: str = "",
    candidates: Optional[list[dict[str, Any]]] = None,
) -> dict[str, Any]:
    blob = text or ""
    state: dict[str, Any] = {"chapter": "", "sublabel": "", "headingBits": {}}
    buckets = _empty_row_buckets()
    change_seen = 0
    parts = re.split(r"(?=【Word )", blob) if "【Word " in blob else [blob]
    if len(parts) == 1:
        # 无训练摘录标记时，按空行切，尽量保留章节标题
        parts = re.split(r"\n{2,}", blob)
    for part in parts:
        _update_heading_state(part, state)
        rows = _parse_tsv_block(part)
        if not rows:
            continue
        kind = _classify_archive_table(
            rows,
            chapter=str(state.get("chapter") or ""),
            sublabel=str(state.get("sublabel") or ""),
            change_seen=change_seen,
        )
        if kind in {"changeCe", "changeNmpa"}:
            change_seen += 1
        if kind and not buckets[kind]:
            buckets[kind] = rows

    source_version = extract_version_label(blob, source_file)
    return _payload(
        rows=buckets,
        heading_bits=state.get("headingBits") or {},
        source_file=source_file,
        source_version=source_version,
        source_date=source_date,
        matched_by=matched_by,
        process_hints=_collect_process_hints(blob),
        candidates=candidates,
        message="已从知识库摘录解析归档表" if any(buckets.values()) else "知识库摘录未解析到归档表",
    )


def pick_latest_rule_file(candidates: list[dict[str, Any]]) -> Optional[dict[str, Any]]:
    """candidates: file_name / created_at / source_version(optional) / text_version(optional)."""
    scored: list[tuple[tuple[Any, ...], dict[str, Any]]] = []
    for row in candidates or []:
        name = str(row.get("file_name") or row.get("sourceFile") or "")
        if not is_yy_iw_020_filename(name) and not _DOC_NO.search(name):
            # 仍允许正文已标出版次的候选
            if not row.get("text_version") and not row.get("source_version"):
                continue
        ver = parse_version_tuple(
            str(row.get("text_version") or row.get("source_version") or "")
            or extract_version_label(name)
        )
        created = str(row.get("created_at") or "")
        scored.append(((ver, created, name), row))
    if not scored:
        return None
    scored.sort(key=lambda x: x[0], reverse=True)
    return scored[0][1]


def extract_from_local_rules_dir(rules_dir: str | Path | None = None) -> dict[str, Any]:
    """扫描 docs/rules 下最新 YY-IW-020。"""
    if rules_dir is None:
        here = Path(__file__).resolve()
        candidates = [
            here.parents[2] / "docs" / "rules",
            here.parents[1] / "docs" / "rules",
        ]
        rules_dir = next((p for p in candidates if p.is_dir()), None)
    root = Path(rules_dir) if rules_dir else None
    if root is None or not root.is_dir():
        return _payload(
            rows=_empty_row_buckets(),
            heading_bits={},
            source_file="",
            source_version="",
            matched_by="local_docx",
            message="未找到本地制度目录",
        )
    files = [
        p
        for p in root.glob("*.docx")
        if is_yy_iw_020_filename(p.name) and not p.name.startswith("~")
    ]
    if not files:
        return _payload(
            rows=_empty_row_buckets(),
            heading_bits={},
            source_file="",
            source_version="",
            matched_by="local_docx",
            message="本地制度目录无 YY-IW-020",
        )
    parsed = [extract_from_docx(p, matched_by="local_docx") for p in files]
    parsed = [x for x in parsed if x.get("ok")]
    if not parsed:
        return extract_from_docx(files[0], matched_by="local_docx")
    parsed.sort(
        key=lambda x: (
            parse_version_tuple(str(x.get("sourceVersion") or "")),
            str(x.get("sourceDate") or ""),
            str(x.get("sourceFile") or ""),
        ),
        reverse=True,
    )
    best = parsed[0]
    best["candidates"] = [
        {
            "file_name": x.get("sourceFile"),
            "source_version": x.get("sourceVersion"),
            "source_date": x.get("sourceDate"),
        }
        for x in parsed
    ]
    return best
