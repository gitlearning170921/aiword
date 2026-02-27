import re
from pathlib import Path
from tkinter import Tk, Label, Entry, Button, messagebox, filedialog, StringVar, Frame, Canvas, Scrollbar, Text

from docx import Document


def replace_placeholders_in_paragraph(paragraph, mapping: dict):
    """
    在一个段落中，用 mapping 中的键值对替换占位符。
    占位符格式：{{key}}
    支持跨 run 的占位符替换
    """
    if not mapping:
        return

    # 方法1：先尝试使用 paragraph.text（会合并所有 run）
    original_text = paragraph.text
    if not original_text:
        return
    
    # 检查是否有占位符
    has_placeholder = False
    for key in mapping.keys():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in original_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return
    
    # 合并所有 run 的文本，确保能捕获跨 run 的占位符
    combined_text = "".join(run.text for run in paragraph.runs if run.text)
    if not combined_text:
        combined_text = original_text
    
    # 替换占位符
    replaced_text = combined_text
    for key, value in mapping.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in replaced_text:
            replaced_text = replaced_text.replace(placeholder, str(value))
    
    # 如果文本有变化，更新段落
    if replaced_text != combined_text:
        # 清空所有 run 并添加新文本
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = replaced_text
        else:
            paragraph.add_run(replaced_text)


def replace_placeholders_in_table(table, mapping: dict):
    """在表格的所有单元格中替换占位符"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, mapping)


def generate_doc_from_template(template_path: str, output_path: str, data: dict):
    """
    :param template_path: 模板 docx 路径，例如 'template.docx'
    :param output_path: 生成的 docx 路径，例如 'output.docx'
    :param data: 占位符映射字典，例如 {"name": "张三", "date": "2025-11-20"}
    """
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    doc = Document(str(template_path))

    # 1. 处理主文档段落
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)

    # 2. 处理主文档表格
    for table in doc.tables:
        replace_placeholders_in_table(table, data)

    # 3. 处理页眉页脚
    for section in doc.sections:
        # 处理页眉
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.header.tables:
                replace_placeholders_in_table(table, data)
        
        # 处理页脚
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.footer.tables:
                replace_placeholders_in_table(table, data)

    # 保存为新文件
    doc.save(output_path)
    print(f"文档已生成: {output_path}")


def extract_placeholders(template_path: str) -> list[str]:
    """从模板文件中提取占位符列表（包括页眉页脚、所有段落、表格等）"""
    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    doc = Document(str(template_path))
    placeholders = []  # 使用列表保持顺序
    seen = set()  # 用于去重
    # 更严格的正则：确保匹配完整的 {{...}}，排除包含换行符或特殊字符的情况
    # 匹配 {{ 和 }} 之间的内容，允许空格，但不允许换行符
    pattern = re.compile(r"\{\{\s*([^{}\n\r]+?)\s*\}\}")

    def collect_from_text(text: str):
        """从文本中提取占位符"""
        if not text or not text.strip():
            return
        
        # 跳过目录行（通常包含制表符和页码）
        if '\t' in text and any(char.isdigit() for char in text):
            return
        
        # 查找所有匹配的占位符
        for match in pattern.finditer(text):
            placeholder = match.group(1).strip()
            # 过滤掉空字符串和只包含空白字符的
            if placeholder and placeholder.strip():
                # 确保占位符不包含大括号（防止误匹配）
                if '{' not in placeholder and '}' not in placeholder:
                    # 保持顺序，只添加未出现过的占位符
                    if placeholder not in seen:
                        placeholders.append(placeholder)
                        seen.add(placeholder)

    def collect_from_paragraph(paragraph):
        """从段落中提取占位符（包括所有 run）"""
        # 合并所有 run 的文本，确保能捕获跨 run 的占位符
        run_texts = []
        for run in paragraph.runs:
            if run.text:
                run_texts.append(run.text)
        
        if run_texts:
            combined_text = "".join(run_texts)
            collect_from_text(combined_text)
        else:
            # 如果没有 run，尝试使用 paragraph.text
            collect_from_text(paragraph.text)

    def collect_from_table(table):
        """从表格中提取占位符（递归处理嵌套表格）"""
        for row in table.rows:
            for cell in row.cells:
                # 处理单元格中的段落
                for paragraph in cell.paragraphs:
                    collect_from_paragraph(paragraph)
                # 处理单元格中的嵌套表格
                for nested_table in cell.tables:
                    collect_from_table(nested_table)

    # 按照文档元素的真实顺序处理主文档（段落和表格交替出现）
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    
    # 创建元素到对象的映射以提高效率
    paragraph_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    
    body = doc.element.body
    for element in body:
        if isinstance(element, CT_P):
            # 这是一个段落
            paragraph = paragraph_map.get(element)
            if paragraph:
                collect_from_paragraph(paragraph)
        elif isinstance(element, CT_Tbl):
            # 这是一个表格
            table = table_map.get(element)
            if table:
                collect_from_table(table)

    # 3. 处理页眉
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                collect_from_paragraph(paragraph)
            # 页眉中的表格
            for table in section.header.tables:
                collect_from_table(table)
        
        # 4. 处理页脚
        if section.footer:
            for paragraph in section.footer.paragraphs:
                collect_from_paragraph(paragraph)
            # 页脚中的表格
            for table in section.footer.tables:
                collect_from_table(table)

    return placeholders


def input_via_dialogs():
    """通过单一弹窗输入所有数据"""
    root = Tk()
    root.title("文档生成工具")
    root.geometry("500x400")
    root.resizable(True, True)
    root.minsize(500, 400)
    
    # 居中显示窗口
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")

    template_path_var = StringVar()
    output_path_var = StringVar()
    placeholder_vars: dict[str, Text] = {}  # 改为存储 Text 组件
    canvas = None
    scrollbar = None
    inner_frame = None

    def clear_placeholder_fields():
        if inner_frame:
            for widget in inner_frame.winfo_children():
                widget.destroy()
        placeholder_vars.clear()

    def build_placeholder_fields(placeholders: list[str]):
        clear_placeholder_fields()
        if not inner_frame:
            return
        if not placeholders:
            Label(
                inner_frame,
                text="模板中未检测到 {{placeholder}} 占位符，请确认模板内容。",
                font=("微软雅黑", 10),
                fg="#888888",
                wraplength=340,
                justify="left",
            ).pack(anchor="w", padx=5, pady=5)
            # 更新滚动区域
            if canvas:
                canvas.update_idletasks()
                canvas.config(scrollregion=canvas.bbox("all"))
            return

        Label(
            inner_frame,
            text="占位符内容：",
            font=("微软雅黑", 10, "bold"),
        ).pack(anchor="w", padx=5, pady=(0, 5))

        for placeholder in placeholders:
            # 创建标签和输入框的容器
            item_frame = Frame(inner_frame)
            item_frame.pack(fill="x", pady=4)
            
            # 标签
            Label(item_frame, text=f"{placeholder}:", width=12, anchor="w", font=("微软雅黑", 9)).pack(side="left", anchor="n", padx=(0, 5))
            
            # 多行文本输入框，支持回车等特殊字符
            text_widget = Text(item_frame, font=("微软雅黑", 9), height=3, wrap="word", undo=True)
            text_widget.pack(side="left", fill="x", expand=True)
            placeholder_vars[placeholder] = text_widget
        
        # 更新滚动区域
        if canvas:
            canvas.update_idletasks()
            canvas.config(scrollregion=canvas.bbox("all"))

    def load_placeholders():
        path = template_path_var.get().strip()
        if not path:
            messagebox.showwarning("警告", "请先选择模板文件")
            return
        try:
            placeholders = extract_placeholders(path)
            build_placeholder_fields(placeholders)
        except FileNotFoundError as exc:
            messagebox.showerror("错误", str(exc))
        except Exception as exc:
            messagebox.showerror("错误", f"解析模板失败：\n{exc}")

    row = 0

    Label(root, text="模板文件:", font=("微软雅黑", 10)).grid(row=row, column=0, padx=15, pady=10, sticky="w")
    frame1 = Frame(root)
    frame1.grid(row=row, column=1, padx=10, pady=10, sticky="ew")
    Entry(frame1, textvariable=template_path_var, width=25, font=("微软雅黑", 9)).pack(side="left", fill="x", expand=True)

    def select_template():
        path = filedialog.askopenfilename(
            title="选择模板文件", filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            template_path_var.set(path)
            load_placeholders()

    Button(frame1, text="浏览", command=select_template, font=("微软雅黑", 9), width=8).pack(side="right", padx=(5, 0))
    Button(frame1, text="载入占位符", command=load_placeholders, font=("微软雅黑", 9), width=10).pack(side="right", padx=(5, 5))
    row += 1

    Label(root, text="输出文件:", font=("微软雅黑", 10)).grid(row=row, column=0, padx=15, pady=10, sticky="w")
    frame2 = Frame(root)
    frame2.grid(row=row, column=1, padx=10, pady=10, sticky="ew")
    Entry(frame2, textvariable=output_path_var, width=25, font=("微软雅黑", 9)).pack(side="left", fill="x", expand=True)

    def select_output():
        path = filedialog.asksaveasfilename(
            title="保存输出文件", defaultextension=".docx", filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if path:
            output_path_var.set(path)

    Button(frame2, text="浏览", command=select_output, font=("微软雅黑", 9), width=8).pack(side="right", padx=(5, 0))
    row += 1

    Label(root, text="占位符输入：", font=("微软雅黑", 10)).grid(row=row, column=0, padx=15, pady=5, sticky="nw")
    
    # 创建带滚动条的占位符输入区域
    placeholders_container = Frame(root, borderwidth=1, relief="groove")
    placeholders_container.grid(row=row, column=1, padx=10, pady=5, sticky="nsew")
    
    # 创建 Canvas 和 Scrollbar
    canvas = Canvas(placeholders_container, highlightthickness=0)
    scrollbar = Scrollbar(placeholders_container, orient="vertical", command=canvas.yview)
    inner_frame = Frame(canvas)
    
    # 将 inner_frame 放入 canvas
    canvas_window = canvas.create_window((0, 0), window=inner_frame, anchor="nw")
    
    # 配置滚动
    def configure_scroll_region(event):
        canvas.update_idletasks()
        canvas.config(scrollregion=canvas.bbox("all"))
    
    def configure_canvas_width(event):
        canvas_width = event.width
        canvas.itemconfig(canvas_window, width=canvas_width)
    
    inner_frame.bind("<Configure>", configure_scroll_region)
    canvas.bind("<Configure>", configure_canvas_width)
    
    # 鼠标滚轮支持
    def on_mousewheel(event):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
    
    canvas.bind_all("<MouseWheel>", on_mousewheel)
    
    # 布局 Canvas 和 Scrollbar
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    canvas.config(yscrollcommand=scrollbar.set)
    
    row += 1

    root.columnconfigure(1, weight=1)
    root.rowconfigure(row - 1, weight=1)

    def on_generate():
        """生成文档按钮点击事件"""
        template_path = template_path_var.get().strip()
        output_path = output_path_var.get().strip()

        if not placeholder_vars:
            messagebox.showwarning("警告", "请先加载并填写模板占位符内容")
            return

        if not template_path:
            messagebox.showwarning("警告", "请选择模板文件")
            return

        if not output_path:
            messagebox.showwarning("警告", "请指定输出文件路径")
            return

        if not Path(template_path).exists():
            messagebox.showerror("错误", f"模板文件不存在: {template_path}")
            return

        missing = [key for key, text_widget in placeholder_vars.items() if not text_widget.get("1.0", "end-1c").strip()]
        if missing:
            messagebox.showwarning("警告", f"以下占位符尚未填写：{', '.join(missing)}")
            return

        try:
            # 从 Text 组件获取内容，去除末尾的换行符
            data = {key: text_widget.get("1.0", "end-1c") for key, text_widget in placeholder_vars.items()}

            generate_doc_from_template(
                template_path=template_path,
                output_path=output_path,
                data=data,
            )

            messagebox.showinfo("成功", f"文档已成功生成！\n输出路径: {output_path}")
            root.destroy()
        except Exception as e:
            messagebox.showerror("错误", f"生成文档时出错:\n{str(e)}")

    # 生成按钮
    Button(root, text="生成文档", command=on_generate, font=("微软雅黑", 12),
           bg="#4CAF50", fg="white", width=15, height=1).grid(row=row, column=0, columnspan=2, pady=20)

    root.mainloop()


if __name__ == "__main__":
    input_via_dialogs()