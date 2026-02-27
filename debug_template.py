"""调试脚本：详细分析 template.docx 的内容"""
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re

def extract_text_from_paragraph(paragraph):
    """提取段落中的所有文本（包括所有 run）"""
    texts = []
    for run in paragraph.runs:
        if run.text:
            texts.append(run.text)
    return "".join(texts)

def analyze_template(template_path):
    """详细分析模板文件"""
    doc = Document(template_path)
    
    print("=" * 60)
    print("分析模板文件内容")
    print("=" * 60)
    
    # 分析主文档段落
    print("\n【主文档段落】")
    for i, para in enumerate(doc.paragraphs, 1):
        text = extract_text_from_paragraph(para)
        if text.strip():
            print(f"段落 {i}: {repr(text)}")
            # 查找占位符
            matches = re.findall(r"\{\{([^}]+)\}\}", text)
            if matches:
                print(f"  找到占位符: {matches}")
    
    # 分析表格
    print("\n【主文档表格】")
    for table_idx, table in enumerate(doc.tables, 1):
        print(f"\n表格 {table_idx}:")
        for row_idx, row in enumerate(table.rows, 1):
            for col_idx, cell in enumerate(row.cells, 1):
                for para_idx, para in enumerate(cell.paragraphs, 1):
                    text = extract_text_from_paragraph(para)
                    if text.strip():
                        print(f"  表格[{row_idx},{col_idx}]段落{para_idx}: {repr(text)}")
                        matches = re.findall(r"\{\{([^}]+)\}\}", text)
                        if matches:
                            print(f"    找到占位符: {matches}")
    
    # 分析页眉页脚
    print("\n【页眉页脚】")
    for section_idx, section in enumerate(doc.sections, 1):
        print(f"\n节 {section_idx}:")
        
        if section.header:
            print("  页眉:")
            for para_idx, para in enumerate(section.header.paragraphs, 1):
                text = extract_text_from_paragraph(para)
                if text.strip():
                    print(f"    段落{para_idx}: {repr(text)}")
                    matches = re.findall(r"\{\{([^}]+)\}\}", text)
                    if matches:
                        print(f"      找到占位符: {matches}")
            for table_idx, table in enumerate(section.header.tables, 1):
                print(f"    表格{table_idx}:")
                for row_idx, row in enumerate(table.rows, 1):
                    for col_idx, cell in enumerate(row.cells, 1):
                        for para_idx, para in enumerate(cell.paragraphs, 1):
                            text = extract_text_from_paragraph(para)
                            if text.strip():
                                print(f"      表格[{row_idx},{col_idx}]段落{para_idx}: {repr(text)}")
                                matches = re.findall(r"\{\{([^}]+)\}\}", text)
                                if matches:
                                    print(f"        找到占位符: {matches}")
        
        if section.footer:
            print("  页脚:")
            for para_idx, para in enumerate(section.footer.paragraphs, 1):
                text = extract_text_from_paragraph(para)
                if text.strip():
                    print(f"    段落{para_idx}: {repr(text)}")
                    matches = re.findall(r"\{\{([^}]+)\}\}", text)
                    if matches:
                        print(f"      找到占位符: {matches}")

if __name__ == "__main__":
    analyze_template("template.docx")

