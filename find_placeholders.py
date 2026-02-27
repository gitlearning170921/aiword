"""查找文档中所有占位符的实际位置"""
from docx import Document
import re

doc = Document("template.docx")
pattern = re.compile(r"\{\{\s*([^{}\n\r]+?)\s*\}\}")

found = {}

def search_text(text, source):
    if not text:
        return
    for match in pattern.finditer(text):
        placeholder = match.group(1).strip()
        if placeholder and '{' not in placeholder and '}' not in placeholder:
            if placeholder not in found:
                found[placeholder] = []
            found[placeholder].append(source)

# 搜索段落
for i, para in enumerate(doc.paragraphs, 1):
    text = "".join(run.text for run in para.runs if run.text)
    if text:
        search_text(text, f"段落{i}")

# 搜索表格
for table_idx, table in enumerate(doc.tables, 1):
    for row_idx, row in enumerate(table.rows, 1):
        for col_idx, cell in enumerate(row.cells, 1):
            for para_idx, para in enumerate(cell.paragraphs, 1):
                text = "".join(run.text for run in para.runs if run.text)
                if text:
                    search_text(text, f"表格{table_idx}[{row_idx},{col_idx}]段落{para_idx}")

print("找到的所有占位符及其位置：")
print("=" * 70)
for placeholder, locations in sorted(found.items()):
    print(f"\n{{{{ {placeholder} }}}}")
    print(f"  出现位置: {', '.join(locations[:3])}{'...' if len(locations) > 3 else ''}")

